import random
import csv
import pandas as pd
import datetime
from datetime import timedelta
import calendar
from openpyxl import load_workbook
import openpyxl
import copy
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT

class Tframe:    # both functions make fill list of targeted time frame with the dates and the day name of the dates
    def __init__(self, datee):
        self.date = datee
        global week, month
        week = []
        month = ''

    def next_30day(self):# fills the dates of the next 30 days
        for i in range(31):
            new_d = self.date + timedelta(days=i)
            main_keys_days.append(new_d.strftime("%A"))
            new_d = new_d.strftime("%d/%m/%Y")
            new_d = f"{new_d}"
            main_dic.update({new_d: []})
            main_keys.append(new_d)

    def next_month(self):# fills the varibles with the dates from the next month
        global month
        days_of_the_month = calendar.monthrange(self.date.year, self.date.month)[1]
        days_till_end_month = days_of_the_month-self.date.day
        self.date = self.date + timedelta(days=days_till_end_month+1)
        days_of_the_month = calendar.monthrange(self.date.year, self.date.month)[1]
        month = self.date.strftime('%B')
        days_name = []

        for i in range(days_of_the_month):
            new_d = self.date + timedelta(days=i)
            main_keys_days.append(new_d.strftime("%A"))
            if new_d.strftime("%A") == "Friday" or new_d.strftime("%A") == "Saturday":
                week.append(str(int(new_d.strftime("%d"))))
            days_name.append(new_d.strftime("%A"))
            new_d = str(int(new_d.strftime("%d")))
            new_d = f"{new_d}"
            main_dic.update({new_d: []})
            main_keys.append(new_d)




class Eframe:
    def __init__(self,main_di, main_di_v, names_list):
        global month
        self.dic_values = main_di_v.copy()
        self.WK_dic = copy.deepcopy(main_di)
        self.N_dic = copy.deepcopy(main_di)
        self.PM_dic = copy.deepcopy(main_di)
        self.names = names_list.copy()

    def N(self):
        global name, N, index, selection_counts
        dicval = self.dic_values.copy()
        names_list = self.names.copy()

        index = 0  # Start from the first name
        random.shuffle(names_list)
        random.shuffle(dicval)
        random.shuffle(names_list)
        random.shuffle(names_list)


        selection_counts = {name: 0 for name in names_list}  # Initialize counts to zero



        for i in range(len(main_keys)):
            k = 0  # Counter to ensure four unique names are added per key
            while k < 4:
                # Sort names by the number of times they've been assigned, ascending
                sorted_names = sorted(selection_counts, key=selection_counts.get)

                sortlen = len(sorted_names)
                # Attempt to assign the least-used name

                for nam in range(sortlen):
                    sorted_names = sorted(selection_counts, key=selection_counts.get)
                    name = sorted_names[nam]



                    # Check for collisions across `AM_dic` and `PM_dic`
                    if name not in self.N_dic[main_keys[i]] and name not in self.N_dic[main_keys[i-1]]:

                        self.N_dic[main_keys[i]].append(name)  # Add to AM_dic
                        selection_counts[name] += 1  # Increment the count for this name
                        k += 1  # Increment unique count for this key
                        break  # Exit inner loop to move to the next unique position


        names_list = self.names.copy()

        N = self.N_dic.copy()
        for e in range(len(main_keys)):
            empp = [""] * len(names_list)

            key = N[main_keys[e]]
            for i in range(len(names_list)):
                # print(i)
                tel = key

                if names_list[i] in tel:
                    empp[i] = "N"
            N[main_keys[e]] = empp

    def PM(self):
        global name, PM, index
        dicval = self.dic_values.copy()
        names_list = self.names.copy()

        index = 0  # Start from the first name
        random.shuffle(names_list)
        random.shuffle(dicval)
        random.shuffle(names_list)
        random.shuffle(names_list)

          # Initialize counts to zero

        for i in range(len(main_keys)):
            k = 0  # Counter to ensure four unique names are added per key
            while k < 4:

                # Sort names by the number of times they've been assigned, ascending
                sorted_names = sorted(selection_counts, key=selection_counts.get)

                sortlen = len(sorted_names)
                # Attempt to assign the least-used name

                for nam in range(sortlen):
                    sorted_names = sorted(selection_counts, key=selection_counts.get)
                    name = sorted_names[nam]

                    # Check for collisions across `AM_dic` and `PM_dic`
                    if name not in self.PM_dic[main_keys[i]] and name not in self.N_dic[main_keys[i-1]]:
                        self.PM_dic[main_keys[i]].append(name)  # Add to AM_dic
                        selection_counts[name] += 1  # Increment the count for this name
                        k += 1  # Increment unique count for this key
                        break  # Exit inner loop to move to the next unique position


        names_list = self.names.copy()




        PM = self.PM_dic.copy()
        for e in range(len(main_keys)):
            empp = [""] * len(names_list)

            key = PM[main_keys[e]]
            for i in range(len(names_list)):
                #print(i)
                tel = key

                if names_list[i] in tel:
                     empp[i] = "PM"
            PM[main_keys[e]] = empp

    def WK(self):
        global name, WK, index, week
        dicval = self.dic_values.copy()
        names_list = self.names.copy()

        index = 0  # Start from the first name
        random.shuffle(names_list)
        random.shuffle(dicval)
        random.shuffle(names_list)
        random.shuffle(names_list)

          # Initialize counts to zero

        for i in range(len(week)):
            k = 0  # Counter to ensure four unique names are added per key
            while k < 4:

                # Sort names by the number of times they've been assigned, ascending
                sorted_names = sorted(selection_counts, key=selection_counts.get)

                sortlen = len(sorted_names)
                # Attempt to assign the least-used name


                for nam in range(sortlen):
                    sorted_names = sorted(selection_counts, key=selection_counts.get)
                    name = sorted_names[nam]

                    # Check for collisions across `AM_dic` and `PM_dic`
                    if name not in self.PM_dic[week[i]] and name not in self.N_dic[main_keys[int(week[i])-2]]\
                            and name not in self.WK_dic[week[i]] and name not in self.N_dic[main_keys[int(week[i])-1]]:
                        self.WK_dic[week[i]].append(name)  # Add to AM_dic
                        selection_counts[name] += 1  # Increment the count for this name
                        k += 1  # Increment unique count for this key
                        break  # Exit inner loop to move to the next unique position


        names_list = self.names.copy()




        WK = self.WK_dic.copy()
        for e in range(len(main_keys)):
            empp = [""] * len(names_list)

            key = WK[main_keys[e]]
            for i in range(len(names_list)):
                #print(i)
                tel = key

                if names_list[i] in tel:
                     empp[i] = "AM"
            WK[main_keys[e]] = empp





    def print(self):
        data = {
            "names":names_list
        }
        global month

        def merge_dicts(dict1, dict2, dict3):
            merged_dict = {}

            for key in dict1.keys():
                # Get the lists from all three dictionaries
                list1 = dict1[key]
                list2 = dict2[key]
                list3 = dict3[key]

                # Initialize the merged list
                merged_list = []

                # Loop through the items in all three lists
                for item1, item2, item3 in zip(list1, list2, list3):
                    if item1 == item2 == item3:
                        # If all items are the same, keep that item
                        merged_list.append(item1)
                    elif ("N" in [item1, item2, item3] and "PM" in [item1, item2, item3]) or \
                            ("N" in [item1, item2, item3] and "WK" in [item1, item2, item3]) or \
                            ("PM" in [item1, item2, item3] and "WK" in [item1, item2, item3]):
                        # If any combination of AM, PM, and WK appears, mark as collision
                        merged_list.append("collision")
                    elif (item1 in ["N", "PM", "WK"] and item2 == "" and item3 == "") or \
                            (item2 in ["N", "PM", "WK"] and item1 == "" and item3 == "") or \
                            (item3 in ["N", "PM", "WK"] and item1 == "" and item2 == ""):
                        # If only one is AM, PM, or WK and the others are empty, keep the non-empty item
                        merged_list.append(item1 if item1 else item2 if item2 else item3)
                    else:
                        # For any other cases, retain empty or other default values
                        merged_list.append("")

                merged_dict[key] = merged_list

            return merged_dict


        main_keys_days.insert(0,' ')


        data.update(merge_dicts(PM, N, WK))
        df = pd.DataFrame(data)
        df.to_excel('schedule.xlsx', index=False, engine='openpyxl')
        wb = load_workbook('schedule.xlsx')  # Replace with your file name
        ws = wb.active
        ws.insert_rows(1)
        for col_num, value in enumerate(main_keys_days, start=1):  # start=1 to start from column A
            ws.cell(row=1, column=col_num, value=value[0])
        ws.column_dimensions['A'].width = lenn+3

        wb.save(f'{month} schedule.xlsx')

        # Create a new Document
        doc = Document()

        # Set the page size to A3 and orientation to Landscape
        section = doc.sections[0]
        section.page_width = Pt(1169)  # A3 width in points (297mm)
        section.page_height = Pt(827)  # A3 height in points (420mm)
        section.orientation = WD_ORIENT.LANDSCAPE  # Set to Landscape orientation
        excel_path = f'{month} schedule.xlsx'
        excel_file = excel_path # Update with your Excel file path
        workbook = openpyxl.load_workbook(excel_file)
        sheet = workbook.active

        # Create a new Word document
        doc = Document()

        # Set the page size to A3 and orientation to Landscape
        section = doc.sections[0]
        section.page_width = Pt(1169)  # A3 width in points
        section.page_height = Pt(827)  # A3 height in points
        section.orientation = WD_ORIENT.LANDSCAPE  # Set to Landscape orientation

        # Create a table based on Excel data
        table = doc.add_table(rows=1, cols=sheet.max_column)  # Initial row (header)

        # Add headers (optional step if your Excel has headers)
        for col in range(sheet.max_column):
            table.cell(0, col).text = str(sheet.cell(row=1, column=col + 1).value)

        # Populate the table with data from the Excel file
        for row_idx in range(2, sheet.max_row + 1):  # Start from row 2 (data rows)
            row = table.add_row()  # Add a new row
            for col_idx in range(1, sheet.max_column + 1):  # Iterate through columns
                cell_value = sheet.cell(row=row_idx, column=col_idx).value
                # Check if the cell is None and replace with an empty string if so
                row.cells[col_idx - 1].text = '' if cell_value is None else str(cell_value)

        # Apply clear borders
        table.style = 'Table Grid'

        # Save the Word document
        doc.save(f'{month} schedule.docx')


        # Save the document






    def count_shifts(self):
        for wq in range(len(self.names)):
            namess = self.names[wq]
            s = 0
            for d in range(len(self.dic_values)):
                rand = self.dic_values[d]
                if namess in self.PM_dic[rand]:
                    s += 1
                if namess in self.N_dic[rand]:
                    s += 1
                if namess in self.WK_dic[rand]:
                    s += 1




            if s == 9 or s==6:
                print(f"{namess}={s}erororororororororooror")
            else:
                print(f"{namess}={s}")

    def count_days_shifts(self):
        for ein in range(len(self.dic_values)):
            l = len(self.N_dic[self.dic_values[ein]])+len(self.PM_dic[self.dic_values[ein]])
            ll =  l +len(self.WK_dic[self.dic_values[ein]])

            print(f"{self.dic_values[ein]}={ll}")



    def count_emps(self):
        namelist = self.names.copy()
        print(f"we have {len(namelist)} available employees")


main_dic = {}  # each key is a date and contains the names of people who are working on these days
main_keys = []  # a list of each key easier to handle
main_keys_days = [] # names of the dates like m for monday
names_list = [] # to be filled with names of the csv file to handle it easier


with open('names.csv', 'r') as na:# fills names in a list
    names = csv.reader(na)
    next(names)# skips the first line in the file
    lenn = 0
    for lk in names:
        if len(lk[0]) > lenn:
            lenn = len(lk[0])
        names_list.append(lk[0])
date = datetime.datetime.now()
for g in range(len(names_list)):
    if len(names_list[g])<lenn:
        names_list[g] = names_list[g]+(''*(lenn-len(names_list[g])))
T = Tframe(date)  # T will fill the necessary list to be able to distribute emps shifts
T.next_month()
e = Eframe(main_dic, main_keys, names_list)

e.N()
e.PM()
e.WK()
e.print()

e.count_emps()
e.count_shifts()
e.count_days_shifts()
